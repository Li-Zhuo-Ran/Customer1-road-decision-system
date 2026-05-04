
import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from docx import Document
from io import BytesIO
import time
import os
import xgboost as xgb
from sklearn.model_selection import KFold
from sklearn.metrics import mean_squared_error
import math

# --- 页面全局配置 ---
st.set_page_config(page_title="路面养护决策系统", layout="wide", initial_sidebar_state="expanded")

# --- 专家知识库: 广西对策库示例 ---
EXPERT_KNOWLEDGE = {
    "雾封层": {
        "场景": "轻度龟裂、表层松散、抗滑性能轻度衰减。PCI≥85，无结构性病害。",
        "参数": "沥青用量0.3~0.5kg/m²；施工环境温度≥10℃。采用广西本地生产的高黏改性乳化沥青。",
        "效果": "PCI提升5~8，预期使用寿命5~6年。",
        "窗口期": "10月~次年4月"
    },
    "同步碎石封层": {
        "场景": "轻度车辙(RD≤5mm)、路面渗水、轻度龟裂。PCI≥80, PSSI≥80。",
        "参数": "沥青用量1.2~1.5kg/m²，碎石撒布覆盖率100%~120%。采用SBS改性沥青(软化点≥60℃)。",
        "效果": "PCI提升8~12，RD降低2~3mm，使用寿命6~8年。",
        "窗口期": "10月~次年3月"
    },
    "局部铣刨重铺": {
        "场景": "中度龟裂、松散、坑槽。70≤PCI<80, PSSI≥75。",
        "参数": "铣刨深度2~4cm，压实度≥96%。采用AC-13C改性沥青混合料。马歇尔稳定度≥10kN。",
        "效果": "PCI提升15~20，病害完全修复，使用寿命8~10年。",
        "窗口期": "10月~次年4月"
    },
    "结构大修": {
        "场景": "PCI<70 或结构强度严重不足。",
        "参数": "全深层铣刨重新铺筑，基层加固。",
        "效果": "彻底解决结构性病害，恢复路面全寿命周期。",
        "窗口期": "全年(避开雨季)"
    }
}

# --- 模拟 PSO-XGBoost 模型核心算法逻辑 ---
def forecast_pci_pssi(df, years=[0,1,2,3,5]):
    """
    按年份预测 PCI 和 PSSI 的简单模拟函数。
    返回一个扩展后的 DataFrame，包含每个路段每年预测的 `预测PCI` 和 `预测PSSI`。
    预测逻辑：基于初始 PCI/PSSI，加上一个由交通量和当前路龄决定的年降幅。
    """
    rows = []
    for _, r in df.iterrows():
        base_pci = float(r.get('PCI', 0))
        base_pssi = float(r.get('PSSI', 0))
        traffic = float(r.get('交通量', 0))
        age = float(r.get('路龄', 0))

        # 年降幅基准（可根据实际模型调整）
        # 交通量越大、路龄越长，降幅越高
        annual_degrade_pci = 1.5 + (traffic / 30000.0) * 6.0 + (age / 20.0) * 2.0
        annual_degrade_pssi = 1.0 + (traffic / 40000.0) * 4.0 + (age / 25.0) * 1.5

        for y in years:
            pred_pci = max(0.0, base_pci - y * annual_degrade_pci)
            pred_pssi = max(0.0, base_pssi - y * annual_degrade_pssi)
            rows.append({
                '路段ID': r.get('路段ID'),
                '年': int(y),
                '预测PCI': round(pred_pci, 2),
                '预测PSSI': round(pred_pssi, 2),
                '交通量': traffic,
                '路龄': age
            })

    return pd.DataFrame(rows)


def compute_pci_pssi_from_measurements(df1, df2):
    """
    将两个按序号对齐的检测表合并并计算 PCI 与 PSSI 的估计值。
    df1: 平整度和构造深度（含 左/右平整度、构造深度、跳车 等）
    df2: 车辙（含 左/右车辙）

    返回：DataFrame，包含列 ['路段ID', 'PCI', 'PSSI', '交通量', '路龄']（若无交通量/路龄则填默认）
    计算规则（简单启发式）：
      - 平整度、跳车、车辙 越大，PCI 越低；归一化为 0-100。
      - 车辙与构造深度影响 PSSI（强度），车辙越深/构造深度越大表示承载能力下降。
    这些是示例启发式公式，建议替换为真实模型。
    """
    # 选择用于合并的键：优先使用 序号 或 起止桩号组合；如果有 '序号' 则用之
    left = df1.copy()
    right = df2.copy()

    if '序号' in left.columns and '序号' in right.columns:
        merged = left.merge(right, on='序号', suffixes=('_p', '_r'))
    elif '起点桩号' in left.columns and '起点桩号' in right.columns:
        merged = left.merge(right, on=['起点桩号','终止桩号'], suffixes=('_p', '_r'))
    else:
        # 兜底：按行索引对齐
        merged = pd.concat([left.reset_index(drop=True), right.reset_index(drop=True)], axis=1)

    # 构造一些指标
    def safe_mean(vals):
        vals = [float(v) for v in vals if pd.notna(v)]
        return sum(vals)/len(vals) if vals else 0.0

    rows = []
    for _, r in merged.iterrows():
        # 平整度参考左/右平整度列名可能为 左平整度/右平整度 或 left/right 平整度
        lp = r.get('左平整度', r.get('left_flatness', np.nan))
        rp = r.get('右平整度', r.get('right_flatness', np.nan))
        left_jump = r.get('左路面跳车', r.get('left_jump', 0))
        right_jump = r.get('右路面跳车', r.get('right_jump', 0))
        rut_l = r.get('左车辙', r.get('左车辙', r.get('left_rut', np.nan)))
        rut_r = r.get('右车辙', r.get('右车辙', r.get('right_rut', np.nan)))
        struct_left = r.get('左构造深度', r.get('left_profile', np.nan))
        struct_right = r.get('右构造深度', r.get('right_profile', np.nan))

        mean_flat = safe_mean([lp, rp])
        mean_jump = safe_mean([left_jump, right_jump])
        mean_rut = safe_mean([rut_l, rut_r])
        mean_struct = safe_mean([struct_left, struct_right])

        # PCI 基本上与平整度正相关，与跳车和车辙负相关
        # 归一化：假设平整度理想值 ~0-5（示例），跳车/车辙越小越好 -> 将它们映射到扣分
        pci_score = 100.0
        # 平整度较小（好），这里把平整度乘以负面系数
        pci_score -= mean_flat * 2.0
        pci_score -= mean_jump * 0.8
        pci_score -= mean_rut * 1.2
        # 结构深度过大也扣分
        pci_score -= max(0.0, mean_struct - 0.5) * 10.0

        # PSSI 主要受车辙和构造深度影响，越大表明结构问题，PSSI 越低
        pssi_score = 100.0 - mean_rut * 2.5 - max(0.0, mean_struct - 0.5) * 8.0 - mean_jump * 0.5

        pci_score = float(np.clip(pci_score, 0, 100))
        pssi_score = float(np.clip(pssi_score, 0, 100))

        rows.append({
            '路段ID': r.get('序号', r.get('起点桩号', None)) or f"ID_{_}",
            'PCI': round(pci_score,2),
            'PSSI': round(pssi_score,2),
            # 保留交通量、路龄列（若存在），否则设为默认
            '交通量': r.get('交通量', r.get('traffic', 5000)),
            '路龄': r.get('路龄', r.get('age', 5))
        })

    return pd.DataFrame(rows)


def _build_features_from_measurements(df_meas_rut):
    # 输入为已合并的测量表（含左/右平整度、跳车、车辙、构造深度、交通量、路龄）
    rows = []
    for _, r in df_meas_rut.iterrows():
        lp = r.get('左平整度', np.nan)
        rp = r.get('右平整度', np.nan)
        left_jump = r.get('左路面跳车', 0)
        right_jump = r.get('右路面跳车', 0)
        rut_l = r.get('左车辙', np.nan)
        rut_r = r.get('右车辙', np.nan)
        struct_left = r.get('左构造深度', np.nan)
        struct_right = r.get('右构造深度', np.nan)

        mean_flat = np.nanmean([v for v in [lp, rp] if pd.notna(v)]) if any(pd.notna(v) for v in [lp, rp]) else 0.0
        mean_jump = np.nanmean([v for v in [left_jump, right_jump] if pd.notna(v)]) if any(pd.notna(v) for v in [left_jump, right_jump]) else 0.0
        mean_rut = np.nanmean([v for v in [rut_l, rut_r] if pd.notna(v)]) if any(pd.notna(v) for v in [rut_l, rut_r]) else 0.0
        mean_struct = np.nanmean([v for v in [struct_left, struct_right] if pd.notna(v)]) if any(pd.notna(v) for v in [struct_left, struct_right]) else 0.0

        rows.append({
            'mean_flat': mean_flat,
            'mean_jump': mean_jump,
            'mean_rut': mean_rut,
            'mean_struct': mean_struct,
            '交通量': r.get('交通量', 5000),
            '路龄': r.get('路龄', 5)
        })
    return pd.DataFrame(rows)


def simple_pso_optimize_xgb(X, y, n_particles=8, n_iters=8, random_state=42):
    """
    简单 PSO 用于 XGBoost 超参（max_depth, learning_rate, n_estimators, subsample）调优。
    返回最佳参数字典。
    """
    np.random.seed(random_state)

    # 参数边界
    bounds = {
        'max_depth': (3, 10),
        'learning_rate': (0.01, 0.5),
        'n_estimators': (50, 500),
        'subsample': (0.5, 1.0)
    }

    # 初始化粒子位置与速度
    particles = []
    velocities = []
    scores = []
    for _ in range(n_particles):
        p = {
            'max_depth': np.random.uniform(*bounds['max_depth']),
            'learning_rate': np.random.uniform(*bounds['learning_rate']),
            'n_estimators': np.random.uniform(*bounds['n_estimators']),
            'subsample': np.random.uniform(*bounds['subsample'])
        }
        v = {k: 0.0 for k in p}
        particles.append(p)
        velocities.append(v)
        scores.append(1e9)

    # 记录个体最优与全局最优
    pbest = particles.copy()
    pbest_scores = scores.copy()
    gbest = None
    gbest_score = 1e9

    def eval_params(param_dict):
        # 转换为 int/float
        params = {
            'max_depth': int(round(param_dict['max_depth'])),
            'learning_rate': float(param_dict['learning_rate']),
            'n_estimators': int(round(param_dict['n_estimators'])),
            'subsample': float(param_dict['subsample'])
        }
        # 5-fold CV RMSE
        kf = KFold(n_splits=3, shuffle=True, random_state=1)
        rmses = []
        for tr_idx, te_idx in kf.split(X):
            Xtr, Xte = X.iloc[tr_idx], X.iloc[te_idx]
            ytr, yte = y.iloc[tr_idx], y.iloc[te_idx]
            model = xgb.XGBRegressor(
                max_depth=params['max_depth'],
                learning_rate=params['learning_rate'],
                n_estimators=params['n_estimators'],
                subsample=params['subsample'],
                objective='reg:squarederror',
                verbosity=0
            )
            model.fit(Xtr, ytr)
            pred = model.predict(Xte)
            rmses.append(math.sqrt(mean_squared_error(yte, pred)))
        return float(np.mean(rmses))

    w = 0.5
    c1 = 0.8
    c2 = 0.9

    for it in range(n_iters):
        for i, p in enumerate(particles):
            score = eval_params(p)
            # 更新个体最优
            if score < pbest_scores[i]:
                pbest[i] = p.copy()
                pbest_scores[i] = score
            # 更新全局最优
            if score < gbest_score:
                gbest = p.copy()
                gbest_score = score

        # 更新速度与位置
        for i in range(n_particles):
            for k in particles[i].keys():
                r1 = np.random.rand()
                r2 = np.random.rand()
                velocities[i][k] = w * velocities[i][k] + c1 * r1 * (pbest[i][k] - particles[i][k]) + c2 * r2 * (gbest[k] - particles[i][k])
                particles[i][k] = particles[i][k] + velocities[i][k]
                # 边界约束
                low, high = bounds[k]
                particles[i][k] = float(np.clip(particles[i][k], low, high))

    best_params = {
        'max_depth': int(round(gbest['max_depth'])),
        'learning_rate': float(gbest['learning_rate']),
        'n_estimators': int(round(gbest['n_estimators'])),
        'subsample': float(gbest['subsample'])
    }
    return best_params, gbest_score

# --- 自定义样式 ---
st.markdown("""
    <style>
    .main { background-color: #f5f7f9; }
    .stButton>button { width: 100%; border-radius: 5px; height: 3em; background-color: #007bff; color: white; }
    .p0-badge { color: red; font-weight: bold; }
    .p1-badge { color: orange; font-weight: bold; }
    </style>
    """, unsafe_allow_html=True)

# --- 初始化 Session State (存储跨页面数据) ---
if 'raw_data' not in st.session_state:
    st.session_state.raw_data = None
if 'processed_data' not in st.session_state:
    st.session_state.processed_data = None

# --- 侧边栏导航 ---
st.sidebar.title("🛣️ 路面养护决策系统")
st.sidebar.info("遵循：数据层 → 模型层 → 规则层 → 优化层 → 输出层")
menu = st.sidebar.radio("核心功能模块", 
    ["M1 数据管理模块 (P0)", 
     "M2 路面性能预测模块 (P0)", 
     "M3 养护决策模块 (P0)", 
     "M4 方案生成模块 (P1)", 
     "M5 效果评估模块 (P1)"])

# --- M1 数据管理模块 ---
if menu == "M1 数据管理模块 (P0)":
    st.markdown("## 📂 M1 数据管理模块 <span class='p0-badge'>[P0]</span>", unsafe_allow_html=True)
    st.subheader("多源数据接入、处理与存储")
    col1, col2 = st.columns([1, 2])
    with col1:
        st.markdown("**上传：平整度与构造深度 (CSV)**")
        file_meas = st.file_uploader("上传 平整度和构造深度 CSV", type=["csv"], key='meas')
        st.markdown("**上传：车辙 (CSV)**")
        file_rut = st.file_uploader("上传 车辙 CSV", type=["csv"], key='rut')

        if st.button("生成示例测试数据（测量表 + 车辙表）"):
            # 生成示例的测量表与车辙表，并计算出 PCI/PSSI
            demo_meas = pd.DataFrame({
                '序号': [i for i in range(1,11)],
                '起点桩号': [179080 + 10*(i-1) for i in range(1,11)],
                '终止桩号': [179080 + 10*i for i in range(1,11)],
                '左平整度': [20.49,2.11,1.73,2.48,3.91,2.59,1.68,1.49,2.08,2.06],
                '右平整度': [19.77,2.68,3.85,2.46,4.84,2.52,1.96,1.85,2.04,2.48],
                '左构造深度': [0.7,0.44,0.45,0.48,0.52,0.41,0.39,0.46,0.45,0.48],
                '右构造深度': [1.09,0.67,0.61,1.03,0.75,0.47,0.43,0.52,0.52,0.45],
                '中构造深度': [1.01,0.48,0.44,0.7,0.63,0.36,0.37,0.41,0.34,0.35],
                '左路面跳车': [146.84,13.18,7.24,9.94,9.51,13.17,4.83,7.38,5.95,6.12],
                '右路面跳车': [61.91,8.87,13.28,4.43,18.38,7.34,3.95,9.71,7.6,8.18],
                '车速': [10.1,13,15.8,19.1,23,24.8,28.1,29.5,30.6,31.7]
            })
            demo_rut = pd.DataFrame({
                '序号': [i for i in range(1,11)],
                '起点桩号': demo_meas['起点桩号'],
                '终止桩号': demo_meas['终止桩号'],
                '左车辙': [15.19,2.21,4.92,2.62,1.56,3.35,1.66,1.43,1.53,1.61],
                '右车辙': [18.61,3.55,5.45,4.97,4.09,3.88,3.3,3.27,4.23,3.65]
            })
            merged_demo = compute_pci_pssi_from_measurements(demo_meas, demo_rut)
            # 保存到 session 以便预览
            st.session_state.meas_raw = demo_meas
            st.session_state.rut_raw = demo_rut
            st.session_state.raw_data = merged_demo
            st.success("示例测量表与车辙表已生成并计算 PCI/PSSI！")

        if file_meas:
            try:
                df_meas = pd.read_csv(file_meas)
                st.session_state.meas_raw = df_meas
                st.success("平整度/构造深度表已上传")
            except Exception as e:
                st.error(f"平整度表解析失败: {e}")

        if file_rut:
            try:
                df_rut = pd.read_csv(file_rut)
                st.session_state.rut_raw = df_rut
                st.success("车辙表已上传")
            except Exception as e:
                st.error(f"车辙表解析失败: {e}")

        # 若同时存在 meas_raw 与 rut_raw，自动合并计算
        if 'meas_raw' in st.session_state and 'rut_raw' in st.session_state:
            try:
                merged = compute_pci_pssi_from_measurements(st.session_state.meas_raw, st.session_state.rut_raw)
                st.session_state.raw_data = merged
                st.success("已基于上传两表计算 PCI/PSSI 并生成预览")
            except Exception as e:
                st.error(f"合并计算失败: {e}")
        elif 'meas_raw' in st.session_state and 'rut_raw' not in st.session_state:
            # 仅测量表，尝试计算但告警
            try:
                merged = compute_pci_pssi_from_measurements(st.session_state.meas_raw, pd.DataFrame())
                st.session_state.raw_data = merged
                st.warning('仅上传测量表，已基于测量表计算 PCI/PSSI（可能不完整）')
            except Exception as e:
                st.error(f"基于测量表计算失败: {e}")

    with col2:
        st.write("### 实时数据预览")
        if 'meas_raw' in st.session_state:
            st.write("**平整度与构造深度 表（预览）**")
            st.dataframe(st.session_state.meas_raw, use_container_width=True)
        if 'rut_raw' in st.session_state:
            st.write("**车辙 表（预览）**")
            st.dataframe(st.session_state.rut_raw, use_container_width=True)
        if st.session_state.raw_data is not None:
            st.write("**计算得到的 PCI / PSSI 结果（预览）**")
            st.dataframe(st.session_state.raw_data, use_container_width=True)

# --- M2 路面性能预测模块 ---
elif menu == "M2 路面性能预测模块 (P0)":
    st.markdown("## 📈 M2 路面性能预测模块 <span class='p0-badge'>[P0]</span>", unsafe_allow_html=True)
    st.subheader("按年份预测 PCI / PSSI 并可视化")

    if st.session_state.raw_data is None:
        st.warning("⚠️ 请先在 M1 模块上传或生成数据")
    else:
        # 年份选择（默认 0,1,2,3,5 年）
        years_input = st.text_input("请输入要预测的年份（逗号分隔，例如：0,1,2,3,4,5,6,7,8,9,10）", value="0,1,2,3,4,5,6,7,8,9,10")
        try:
            years = [int(x.strip()) for x in years_input.split(',') if x.strip()!='']
        except:
            years = [0,1,2,3,5]

        if st.button("🚀 运行按年预测 (PCI / PSSI)"):
            with st.spinner('正在模拟按年预测...'):
                time.sleep(1.0)
                df = st.session_state.raw_data.copy()
                # 如果 raw_data 是测量表（未计算 PCI/PSSI），尝试用启发式函数先得到基线
                if 'PCI' not in df.columns or 'PSSI' not in df.columns:
                    # 假设上传的是测量表，已在 M1 用 compute_pci_pssi_from_measurements 生成
                    base_df = df.copy()
                else:
                    base_df = df.copy()
                # 如果存在测量相关列，尝试用模型训练（PSO-XGBoost）预测更稳定的基线
                st.session_state.processed_data = None
                st.session_state.pso_models = {}
                st.session_state.pso_results = {}
                st.success("按年预测基础数据已准备（下一步可选择运行 PSO-XGBoost 训练）")

        # PSO-XGBoost 训练区：在已有 raw_data 的基础上训练模型
        if st.button("⚙️ 使用 PSO 优化并训练 XGBoost (适用于大数据)"):
            if st.session_state.raw_data is None:
                st.error('请先在 M1 上传或生成数据')
            else:
                with st.spinner('正在准备训练数据...'):
                    df0 = st.session_state.raw_data.copy()
                    # 若上传的是测量表，df0 可能包含原始测量列；我们需要合并测量以构造特征
                    # 如果数据已经是合并后的特征表（含 mean_*），直接用之
                    feature_df = None
                    if set(['mean_flat','mean_jump','mean_rut','mean_struct']).issubset(df0.columns):
                        feature_df = df0[['mean_flat','mean_jump','mean_rut','mean_struct','交通量','路龄']].copy()
                    else:
                        # 尝试按现有列构建特征
                        feature_df = _build_features_from_measurements(df0)

                    # 目标 y：如果有 PCI/PSSI 列则用原始值，否则用启发式计算的值
                    if 'PCI' in df0.columns and 'PSSI' in df0.columns:
                        y_pci = df0['PCI']
                        y_pssi = df0['PSSI']
                    else:
                        # 使用启发式函数输出
                        tmp = compute_pci_pssi_from_measurements(df0, pd.DataFrame())
                        y_pci = tmp['PCI']
                        y_pssi = tmp['PSSI']

                # 训练 PCI 模型
                with st.spinner('正在用 PSO 优化 XGBoost (PCI)...'):
                    try:
                        best_pci_params, pci_score = simple_pso_optimize_xgb(feature_df, y_pci, n_particles=8, n_iters=6)
                        st.session_state.pso_results['pci'] = (best_pci_params, pci_score)
                        st.success(f"PCI 模型最佳参数: {best_pci_params}  CV RMSE={pci_score:.3f}")
                    except Exception as e:
                        st.error(f"PCI 模型训练失败: {e}")

                # 训练 PSSI 模型
                with st.spinner('正在用 PSO 优化 XGBoost (PSSI)...'):
                    try:
                        best_pssi_params, pssi_score = simple_pso_optimize_xgb(feature_df, y_pssi, n_particles=8, n_iters=6)
                        st.session_state.pso_results['pssi'] = (best_pssi_params, pssi_score)
                        st.success(f"PSSI 模型最佳参数: {best_pssi_params}  CV RMSE={pssi_score:.3f}")
                    except Exception as e:
                        st.error(f"PSSI 模型训练失败: {e}")

                # 用最佳参数训练最终模型并保存
                with st.spinner('训练最终模型并生成年度预测...'):
                    model_pci = xgb.XGBRegressor(**best_pci_params, objective='reg:squarederror', verbosity=0)
                    model_pssi = xgb.XGBRegressor(**best_pssi_params, objective='reg:squarederror', verbosity=0)
                    model_pci.fit(feature_df, y_pci)
                    model_pssi.fit(feature_df, y_pssi)
                    st.session_state.pso_models['pci'] = model_pci
                    st.session_state.pso_models['pssi'] = model_pssi

                    # 生成基线预测并按年展开
                    base_pred = feature_df.copy()
                    base_pred['PCI'] = model_pci.predict(feature_df)
                    base_pred['PSSI'] = model_pssi.predict(feature_df)
                    # 构造 base_df 与原始索引对应的路段ID
                    base_df = pd.DataFrame({
                        '路段ID': df0.get('路段ID', pd.Series([f'R{i}' for i in range(len(base_pred))])),
                        'PCI': base_pred['PCI'].values,
                        'PSSI': base_pred['PSSI'].values,
                        '交通量': base_pred.get('交通量', df0.get('交通量', 5000)),
                        '路龄': base_pred.get('路龄', df0.get('路龄', 5))
                    })

                    forecast_df = forecast_pci_pssi(base_df, years=years)
                    st.session_state.processed_data = forecast_df
                    st.success('PSO-XGBoost 训练并年度预测完成')

        if st.session_state.processed_data is not None:
            forecast_df = st.session_state.processed_data
            available_years = sorted(forecast_df['年'].unique())
            sel_year = st.selectbox("选择显示年份", available_years)
            df_year = forecast_df[forecast_df['年'] == int(sel_year)]

            c1, c2 = st.columns(2)
            with c1:
                if not df_year.empty:
                    df_melt = df_year.melt(id_vars=['路段ID'], value_vars=['预测PCI','预测PSSI'], var_name='指标', value_name='值')
                    fig = px.bar(df_melt, x='路段ID', y='值', color='指标', barmode='group',
                                 title=f'各路段 {sel_year} 年预测 PCI / PSSI')
                    st.plotly_chart(fig)
                else:
                    st.info('所选年份无数据')
            with c2:
                if not df_year.empty:
                    fig2 = px.scatter(df_year, x='交通量', y='预测PCI', size='路龄', hover_name='路段ID',
                                     title=f'交通量 vs 预测PCI (年={sel_year})')
                    st.plotly_chart(fig2)
                else:
                    st.info('所选年份无数据')

# --- M3 养护决策模块 (基于专家知识库) ---
elif menu == "M3 养护决策模块 (P0)":
    st.header("🧠 智能决策生成 ")
    if st.session_state.processed_data is None:
        st.warning("请先在 M2 运行预测")
    else:
        df_all = st.session_state.processed_data.copy()

        # 如果按年展开，选择年份
        if '年' in df_all.columns:
            years_available = sorted(df_all['年'].unique())
            sel_year = st.selectbox('请选择要决策的年份', years_available)
            df = df_all[df_all['年'] == int(sel_year)].copy()
        else:
            sel_year = None
            df = df_all.copy()

        # 先确保存在预测健康分列，否则用 PCI/PSSI 加权生成
        if '预测健康分' not in df.columns:
            df['预测健康分'] = df.apply(lambda x: (float(x.get('预测PCI', x.get('PCI', 0))) * 0.6 + float(x.get('预测PSSI', x.get('PSSI', 0))) * 0.4), axis=1)

        def get_pro_decision(row):
            score = float(row.get('预测健康分', 0))
            if score >= 85:
                return "雾封层"
            elif 80 <= score < 85:
                return "同步碎石封层"
            elif 70 <= score < 80:
                return "局部铣刨重铺"
            else:
                return "结构大修"

        df['决策方案'] = df.apply(get_pro_decision, axis=1)

        # 将专家知识详情合并到表中，便于后续导出
        def attach_expert_info(scheme):
            info = EXPERT_KNOWLEDGE.get(scheme, {})
            return info.get('参数',''), info.get('窗口期',''), info.get('效果','')

        df[['核心参数', '建议施工期', '预期效果']] = df['决策方案'].apply(lambda s: pd.Series(attach_expert_info(s)))

        # 合并回 processed_data
        if '年' in df_all.columns:
            df_all = df_all.merge(df[['路段ID','年','决策方案','核心参数','建议施工期','预期效果','预测健康分']], on=['路段ID','年'], how='left')
        else:
            df_all = df_all.merge(df[['路段ID','决策方案','核心参数','建议施工期','预期效果','预测健康分']], on=['路段ID'], how='left')
        st.session_state.processed_data = df_all

        st.write("### 自动化建议列表")
        display_cols = ['路段ID','预测健康分','决策方案','核心参数','建议施工期']
        available = [c for c in display_cols if c in df.columns]
        st.table(df[available])

# --- M4 方案生成模块 (详细专业版) ---
elif menu == "M4 方案生成模块 (P1)":
    st.header('📄 标准化方案导出')
    if st.session_state.processed_data is None or '决策方案' not in st.session_state.processed_data.columns:
        st.error("请先在 M3 生成决策")
    else:
        df_all = st.session_state.processed_data
        # 年份支持
        if '年' in df_all.columns:
            years_available = sorted(df_all['年'].unique())
            sel_year = st.selectbox('选择要导出的年份', years_available)
            df_year = df_all[df_all['年'] == int(sel_year)].copy()
        else:
            sel_year = None
            df_year = df_all.copy()

        if not df_year.empty:
            ids = df_year['路段ID'].unique().tolist()
            selected_id = st.selectbox('选择路段', ids)
            row = df_year[df_year['路段ID'] == selected_id].iloc[0]

            scheme = row.get('决策方案', '')
            detail = EXPERT_KNOWLEDGE.get(scheme, {})

            # 回退到 processed_data 中已附带的字段（如果专家库中缺失）
            core_param = detail.get('参数') or row.get('核心参数') or '未提供核心参数'
            window = detail.get('窗口期') or row.get('建议施工期') or '未提供建议施工期'
            effect = detail.get('效果') or row.get('预期效果') or '未提供预期效果'
            scene = detail.get('场景') or row.get('适用场景') or '未提供适用场景'

            if st.button('生成详细技术方案'):
                # 在界面提示缺失信息来源
                if not detail:
                    st.warning('注意：专家知识库未包含该方案，文档中将使用 M3 中生成的字段或默认文本。')

                doc = Document()
                doc.add_heading(f'路段 {selected_id} 养护技术方案', 0)

                # 一、 路段基本概况
                doc.add_heading('一、 路段基本概况', level=1)
                p = doc.add_paragraph()
                p.add_run('根据系统预测，该路段当前预测健康评分为 ').bold = True
                try:
                    p.add_run(f"{float(row.get('预测健康分', 0)):.2f}分").underline = True
                except Exception:
                    p.add_run(str(row.get('预测健康分', 'N/A')))
                p.add_run('。推荐方案为：')
                p.add_run(f"{scheme}").bold = True

                # 二、 核心技术要求（表格）
                doc.add_heading('二、 核心技术要求', level=1)
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '维度'
                hdr_cells[1].text = '具体要求/指标'

                items = [
                    ("适用场景", scene),
                    ("核心技术参数", core_param),
                    ("预期养护效果", effect),
                    ("建议施工期", window),
                ]
                for item, val in items:
                    row_cells = table.add_row().cells
                    row_cells[0].text = item
                    row_cells[1].text = str(val)

                # 三、 施工关键控制点
                doc.add_heading('三、 施工关键控制点', level=1)
                doc.add_paragraph('1. 施工前需进行现场核查，确认病害范围与系统预测一致。')
                doc.add_paragraph(f'2. {scheme} 作业应严格控制温度及用量，确保施工环境符合规范。')
                doc.add_paragraph('3. 完工后应严格执行封道养护时间，待路面强度达标后方可开放交通。')

                buffer = BytesIO()
                doc.save(buffer)
                st.download_button(
                    label='📥 下载 Word 专业版方案',
                    data=buffer.getvalue(),
                    file_name=f'Road_{selected_id}_Technical_Plan.docx',
                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
        else:
            st.info('所选年份/数据下无可导出的路段')

# --- M5 效果评估模块 ---
elif menu == "M5 效果评估模块 (P1)":
    st.markdown("## 📊 M5 效果评估模块 <span class='p1-badge'>[P1]</span>", unsafe_allow_html=True)
    st.subheader("养护后跟踪评估与反馈")

    st.write("### 养护前后性能对比模拟")
    # 模拟对比数据
    compare_data = pd.DataFrame({
        '阶段': ['养护前', '养护后'],
        'PCI (破损指数)': [65, 98],
        'RQI (平整度)': [60, 95],
        '安全系数': [0.7, 0.95]
    })
    
    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(r=[65, 60, 0.7*100], theta=['PCI','RQI','安全'], fill='toself', name='养护前'))
    fig.add_trace(go.Scatterpolar(r=[98, 95, 0.95*100], theta=['PCI','RQI','安全'], fill='toself', name='养护后'))
    fig.update_layout(polar=dict(radialaxis=dict(visible=True, range=[0, 100])), showlegend=True, title="路段改善雷达图")
    st.plotly_chart(fig)
    st.success("评估结果：养护后综合性能提升了 45%！")
